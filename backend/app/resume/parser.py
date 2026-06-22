"""EMBEDHUNT AI — Resume Parser (PDF/DOCX/TXT)"""
import io
from dataclasses import dataclass
from enum import Enum

class FileType(str, Enum):
    PDF="pdf"; DOCX="docx"; TXT="txt"; UNSUPPORTED="unsupported"

@dataclass
class ParsedDocument:
    raw_text: str; file_type: FileType; page_count: int
    char_count: int; is_scanned: bool; warnings: list[str]

def detect_file_type(filename: str, content: bytes) -> FileType:
    if content[:4] == b"%PDF": return FileType.PDF
    if content[:2] == b"PK" and filename.lower().endswith(".docx"): return FileType.DOCX
    if filename.lower().endswith(".txt"): return FileType.TXT
    if filename.lower().endswith(".pdf"): return FileType.PDF
    return FileType.UNSUPPORTED

def _parse_pdf(content: bytes) -> ParsedDocument:
    warnings, parts, page_count = [], [], 0
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            page_count = len(pdf.pages)
            for p in pdf.pages:
                t = p.extract_text()
                if t: parts.append(t)
    except Exception as e:
        warnings.append(f"Primary parser: {e}")
        try:
            import PyPDF2
            r = PyPDF2.PdfReader(io.BytesIO(content)); page_count = len(r.pages)
            for p in r.pages: parts.append(p.extract_text() or "")
        except Exception as e2: warnings.append(f"Fallback: {e2}")
    text = "\n".join(parts).strip()
    is_scanned = page_count > 0 and len(text) < 100
    if is_scanned: warnings.append("Scanned PDF detected. Text extraction incomplete.")
    return ParsedDocument(text, FileType.PDF, page_count, len(text), is_scanned, warnings)

def _parse_docx(content: bytes) -> ParsedDocument:
    warnings = []
    try:
        import docx; doc = docx.Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    if cell.text.strip(): parts.append(cell.text.strip())
        text = "\n".join(parts)
    except Exception as e: warnings.append(f"DOCX error: {e}"); text = ""
    return ParsedDocument(text, FileType.DOCX, 1, len(text), False, warnings)

def parse_resume(filename: str, content: bytes) -> ParsedDocument:
    ft = detect_file_type(filename, content)
    if ft == FileType.PDF: return _parse_pdf(content)
    if ft == FileType.DOCX: return _parse_docx(content)
    if ft == FileType.TXT: return ParsedDocument(content.decode("utf-8","replace"), FileType.TXT, 1, len(content), False, [])
    return ParsedDocument("", FileType.UNSUPPORTED, 0, 0, False, [f"Unsupported: {filename}"])
